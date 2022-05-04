#!/usr/bin/env python
# coding: utf-8


from tkinter import *
from tkinter.ttk import *
from tkinter import messagebox
from datetime import datetime
import time
import pandas as pd

import docx
from docx.enum.text import WD_COLOR_INDEX


class Compear():
    def __init__(self, string1, string2, Output_name):
        self.str1 = string1.get().upper()
        self.str2 = string2.get().upper()
        self.minnumber = 3  # give by Lorenso
        self.strsplit = []
        self.doc = docx.Document()
        self.AminoTable = {
            "A": [1, 89.1], "G": [1, 75.07], "I": [2, 131.18], "L": [2, 131.18], "V": [2, 117.15],
            "M": [3, 149.21], "F": [4, 165.19], "W": [4, 204.23], "Y": [4, 181.19], "N": [5, 132.12],
            "D": [5, 133.11], "Q": [5, 146.15], "E": [5, 147.13], "C": [6, 121.16], "S": [6, 105.09],
            "T": [6, 119.12], "R": [7, 174.2], "K": [7, 146.19], "H": [8, 155.16], "P": [9, 115.13]
        }

        # Assuming heare everithing is alphabet letters only

        self.str1_n = self.Protein_to_number(self.str1)
        self.str2_n = self.Protein_to_number(self.str2)

        self.Word_document_init(self.doc)

        # split 1st string in the minnumber (3) exmp in "ABCD"
        # out "ABC" and "BCD"
        for i in range(len(self.str1_n) - self.minnumber + 1):
            self.strsplit.append(self.str1_n[i:i + self.minnumber])

        r = 0
        # Iidx and Fidx belong to str2
        df = pd.DataFrame(columns=["Match", "Iidx", "Fidx", "Origiidx", "Origfidx"])

        # here we find all the exact matches and we save them to a df
        for i in range(len(self.strsplit)):
            while (r < len(self.str2_n) - self.minnumber and r != -1):
                r = self.str2_n.find(self.strsplit[i], r)
                if r != -1:
                    df = df.append({"Match": self.strsplit[i],
                                    "Iidx": r, "Fidx": r + self.minnumber - 1, "Origiidx": i,
                                    "Origfidx": i + self.minnumber - 1}, ignore_index=True)
                    r += 1

            r = 0
        print("Untill heart the df hass all the 3 exact matches")
        self.df = df
        print(df)
        df_final = self.Number_to_protein(df, self.str1, self.str2)
        print(df_final)

        self.write_df_to_doc(df_final, self.doc)

        self.mark_strings_in_word(df_final, self.doc)

        # save still in test
        self.Save_word_doc(Output_name, self.doc)

    def extend_final_idx(self):
        for i in range(self.df.shape[0]):
            iidx = self.df.iloc[i][1]
            fidx = self.df.iloc[i][2]  # final index location
            oidx = self.df.iloc[i][3]
            ad = self.minnumber
            # sumand=0
            f = (oidx + 1) + ad  # add one to companeste chacking
            g = (fidx + 2)
            while (f < len(self.str1)):

                if (self.str1[oidx:f] == self.str2[iidx:g]):
                    self.df.loc[i, "Fidx"] = g - 1
                    self.df.loc[i, "Origfidx"] = f - 1
                    self.df.loc[i, "Match"] = self.str1[oidx:f]
                    f += 1
                    if g < len(self.str2):
                        g += 1
                else:
                    break

        print("Matches with extended index are:")
        print(self.df)

    def write_df_to_doc(self, df, doc):
        """
        Write a Data frame to the word document

        Atributes:
            df: dataframe
            doc: word document using docx library
        Return:
            None, changes made to the word document
        """
        t = doc.add_table(df.shape[0] + 1, df.shape[1])
        # add header rows
        for j in range(df.shape[-1]):
            t.cell(0, j).text = df.columns[j]
        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i + 1, j).text = str(df.values[i, j])

    def Protein_to_number(self, string):
        """
        This will conbert the string to a specific grop of proteins for example
            A is similar to G they will both be define as 1

        """
        st = ""
        for i in string:
            st += str(self.AminoTable[i][0])

        return st

    def get_molecular_weight(self, string):
        weight = 0
        for i in string:
            weight += self.AminoTable[i][1]

        return weight

    def Number_to_protein(self, df, st1, st2):
        """
        Combert the input numpy array back to protein gormat givving a clear mattching
        of the proteins with index

        Parameters
        ----------
        df : Data Frame
            this will contain the matching string, initian anf final index of
            the match in both strings
        st1 : String 1
            input for the user
        st2 : String 2
            input for the user
        Returns
        -------
        Df with the all the matches and indexes
        """
        df2 = pd.DataFrame(columns=["String1", "Initial_idx", "Final_idx", "String2", "Initial_idx1", \
                                    "Final_idx1", "Exact_match"])

        for i in range(df.shape[0]):
            idx1 = df.loc[i, "Origiidx"]
            idx2 = df.loc[i, "Origfidx"]
            idx3 = df.loc[i, "Iidx"]
            idx4 = df.loc[i, "Fidx"]

            df2.loc[i, "String1"] = self.str1[idx1:idx2 + 1]
            df2.loc[i, "Initial_idx"] = idx1
            df2.loc[i, "Final_idx"] = idx2
            df2.loc[i, "String2"] = self.str2[idx3:idx4 + 1]
            df2.loc[i, "Initial_idx1"] = idx3
            df2.loc[i, "Final_idx1"] = idx4
            if (self.str1[idx1:idx2 + 1] == self.str2[idx3:idx4 + 1]):  # if st1 == st2 is a exact match so 1
                df2.loc[i, "Exact_match"] = 1
            else:
                df2.loc[i, "Exact_match"] = 0

        return df2

    def Save_word_doc(self, Title, doc):
        doc.save(Title + '.docx')

    def Word_document_init(self, doc):
        # add heading
        doc.add_heading('Comparison Outcome', 0)
        # add paragraph
        para = doc.add_paragraph(
            "Original Sequence: " + self.str1
        )
        mol_weight_str1 = doc.add_paragraph(
            "Molecular Weight of Original Sequence: " + self.get_molecular_weight(self.str1)
        )
        para0 = doc.add_paragraph(
            "Comparison Sequence: " + self.str2
        )
        mol_weight_str2 = doc.add_paragraph(
            "Molecular Weight of Comparison Sequence: " + self.get_molecular_weight(self.str2)
        )
        para1 = doc.add_paragraph(
            "The following sequence will have the identical and similar sequences of the Comparison Sequence"
            "highlighted. SIMILAR SEQUENCES: PINK HIGHLIGHT   IDENTICAL SEQUENCES: TURQUOISE HIGHLIGHT"
        )

    def mark_strings_in_word(self, df, doc):
        """
        Write and higligth the strings in the word document

        Parameters
        ----------
        df : Pandas DataFrame
            Contain the index that march in both strings
        doc : Word Document
            Document where we store all the strinfs
        Returns
        -------
        None.
        """
        # here will need to sort the index
        j: int = 0
        i: int = 0
        #First Protein
        p1 = doc.add_paragraph("Protein #1:")
        sim_para = doc.add_paragraph()

        while i in range(df.shape[0]):
            if j in range(len(self.str1)):
                #print part of the string that does not have and similar or identical sequences
                sim_para.add_run(self.str1[j:df.loc[i, "Initial_idx"]])
                #check for runs that have the same starting index
                if df.loc[i, "Initial_idx"] == df.loc[i+1, "Initial_idx"]:
                    #find the range of the two runs
                    rng_i = df.loc[i, "Final_idx"] - df.loc[i, "Initial_idx"]
                    rng_i2 = df.loc[i+1, "Final_idx"] - df.loc[i+1,"Initial_idx"]
                    # if the 2nd run has a larger range than the first one prioritize it
                    if rng_i < rng_i2:
                        #if it is an exact match highlight with turquoise
                        if df.loc[i+1, "Exact_match"] == 1:
                            sim_para.add_run(
                                self.str1[df.loc[i+1, "Initial_idx"]:df.loc[i,"Final_idx"]]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        #else highlight it in pink
                        else:
                            sim_para.add_run(
                                self.str1[df.loc[i + 1, "Initial_idx"]:df.loc[i, "Final_idx"]]
                            ).font.highlight_color = WD_COLOR_INDEX.PINK
                        #increment j
                        j = df.loc[i + 1, "Final_idx"]
                        #increment i so that next time around i + 1 will be skipped
                        i = i + 1
                    #if the ranges are the same check for which one is an exact match
                    elif rng_i == rng_i2:
                        if df.loc[i, "Exact_match"] == 1:
                            sim_para.add_run(
                                self.str1[df.loc[i,"Initial_idx"]:df.loc[i,"Final_idx"]]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        elif df.loc[i+1, "Exact_match"] == 1:
                            sim_para.add_run(
                                self.str1[df.loc[i+1,"Initial_idx"]:df.loc[i,"Final_idx"]]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        #if they are both just similarities highlight with pink
                        else:
                            sim_para.add_run(
                                self.str1[df.loc[i,"Initial_idx"]:df.loc[i,"Final_idx"]]
                            ).font.highlight_color = WD_COLOR_INDEX.PINK
                        j = df.loc[i, "Final_idx"]
                        i = i + 1
                    #if run 1 has a larger range than run 2
                    else:
                        #if exact match highlight turquoise
                        if df.loc[i, "Exact_match"] == 1:
                            sim_para.add_run(
                                self.str1[df.loc[i, "Initial_idx"]:df.loc[i,"Final_idx"]]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        #else highlight with pink
                        else:
                            sim_para.add_run(
                                self.str1[df.loc[i, "Initial_idx"]:df.loc[i,"Final_idx"]]
                            ).font.highlight_color = WD_COLOR_INDEX.PINK
                        j = df.loc[i, "Final_idx"]
                        i = i + 1
                #if run i's initial index does not match i + 1's index then just
                #check if it is an exact match or not and highlight accordingly
                elif df.loc[i, "Exact_match"] == 1:
                    sim_para.add_run(
                        self.str1[df.loc[i, "Initial_idx"]:df.loc[i, "Final_idx"]]
                    ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    j = df.loc[i, "Final_idx"]
                else:
                    sim_para.add_run(
                        self.str1[df.loc[i, "Initial_idx"]:df.loc[i, "Final_idx"]]
                    ).font.highlight_color = WD_COLOR_INDEX.PINK
                    j = df.loc[i, "Final_idx"]
            i = i + 1
        #clean up
        if j != len(self.str1):
            sim_para.add_run(
                self.str1[j:len(self.str1)]
            )
        #Second protein
        i = 0
        j = 0
        #Sort data frame so index's of second sequence are in order (Ascending)
        df.sort_values(by = "Initial_idx1")
        p2 = doc.add_paragraph("Protein #2:")
        sim_para = doc.add_paragraph()
        while i in range(df.shape[0]):
            if j in range(len(self.str1)):
                #print part of the string           that does not have and similar or identical sequences
                sim_para.add_run(self.str2[j:df.loc[i, "Initial_idx1"]])
                #check for runs that have the same starting index
                if df.loc[i, "Initial_idx1"] == df.loc[i+1, "Initial_idx1"]:
                    #find the range of the two runs
                    rng_i = df.loc[i, "Final_idx1"] - df.loc[i, "Initial_idx1"]
                    rng_i2 = df.loc[i+1, "Final_idx1"] - df.loc[i+1,"Initial_idx1"]
                    # if the 2nd run has a larger range than the first one prioritize it
                    if rng_i < rng_i2:
                        #if it is an exact match highlight with turquoise
                        if df.loc[i+1, "Exact_match"] == 1:
                            sim_para.add_run(
                                self.str2[df.loc[i+1, "Initial_idx1"]:df.loc[i,"Final_idx1"]]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        #else highlight it in pink
                        else:
                            sim_para.add_run(
                                self.str2[df.loc[i + 1, "Initial_idx1"]:df.loc[i, "Final_idx1"]]
                            ).font.highlight_color = WD_COLOR_INDEX.PINK
                        #increment j
                        j = df.loc[i + 1, "Final_idx1"]
                        #increment i so that next time around i + 1 will be skipped
                        i = i + 1
                    #if the ranges are the same check for which one is an exact match
                    elif rng_i == rng_i2:
                        if df.loc[i, "Exact_match"] == 1:
                            sim_para.add_run(
                                self.str2[df.loc[i,"Initial_idx1"]:df.loc[i,"Final_idx1"]]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        elif df.loc[i+1, "Exact_match"] == 1:
                            sim_para.add_run(
                                self.str2[df.loc[i+1,"Initial_idx1"]:df.loc[i,"Final_idx1"]]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        #if they are both just similarities highlight with pink
                        else:
                            sim_para.add_run(
                                self.str2[df.loc[i,"Initial_idx1"]:df.loc[i,"Final_idx1"]]
                            ).font.highlight_color = WD_COLOR_INDEX.PINK
                        j = df.loc[i, "Final_idx1"]
                        i = i + 1
                    #if run 1 has a larger range than run 2
                    else:
                        #if exact match highlight turquoise
                        if df.loc[i, "Exact_match"] == 1:
                            sim_para.add_run(
                                self.str2[df.loc[i, "Initial_idx1"]:df.loc[i,"Final_idx1"]]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                        #else highlight with pink
                        else:
                            sim_para.add_run(
                                self.str2[df.loc[i, "Initial_idx1"]:df.loc[i,"Final_idx1"]]
                            ).font.highlight_color = WD_COLOR_INDEX.PINK
                        j = df.loc[i, "Final_idx1"]
                        i = i + 1
                #if run i's initial index does not match i + 1's index then just
                #check if it is an exact match or not and highlight accordingly
                elif df.loc[i, "Exact_match"] == 1:
                    sim_para.add_run(
                        self.str2[df.loc[i, "Initial_idx1"]:df.loc[i, "Final_idx1"]]
                    ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    j = df.loc[i, "Final_idx1"]
                else:
                    sim_para.add_run(
                        self.str2[df.loc[i, "Initial_idx1"]:df.loc[i, "Final_idx1"]]
                    ).font.highlight_color = WD_COLOR_INDEX.PINK
                    j = df.loc[i, "Final_idx1"]
            i = i + 1
        #clean up
        if j != len(self.str2):
            sim_para.add_run(
                self.str2[j:len(self.str2)]
            )


class GUI():
    def __init__(self):
        """
        GUI is a class that creat the user interface to compate two strings

         Attributes:
        -----------
        None
        Methods:
        --------
        Constructor will crate the gui and the following methods will help to set info and generete outputs
            defname()
            manualname()
            mesagefinish()
            cleanscrean()
            proteincompare()
            chektext()
            startgui()
            printinfo()



        """

        # initialise window
        self.window = Tk()
        self.window.geometry("1000x650")

        # set up labels
        l1 = Label(self.window, text="String 1:")
        l2 = Label(self.window, text="String 2:")
        l3 = Label(self.window, text="Output Name:")
        # place lables
        l1.place(x=20, y=125)
        l2.place(x=20, y=350)
        l3.place(x=20, y=550)

        # define entrys
        # Entery text boxes
        self.String1 = StringVar()
        self.e1 = Entry(self.window, textvariable=self.String1)
        self.e1.place(x=100, y=50, width=800, height=200)
        self.String2 = StringVar()
        self.e2 = Entry(self.window, textvariable=self.String2)
        self.e2.place(x=100, y=275, width=800, height=200)
        self.outText = StringVar()
        self.e3 = Entry(self.window, textvariable=self.outText)
        self.e3.place(x=100, y=550, width=400, height=25)

        ##defauld output name
        self.name = datetime.now().strftime('%Y-%m-%d_%H-%M-%S') + "_output"

        # Output Name set up
        self.v = IntVar()
        self.v.set(1)  # Def name option selected by default
        l4 = Label(self.window, text="Output Name")
        l4.place(x=550, y=500)

        # Radiobuttons
        self.R1 = Radiobutton(self.window, text="Default name", variable=self.v, value=1, command=self.defname)
        self.R1.place(x=550, y=525)
        self.R2 = Radiobutton(self.window, text="Manual name", variable=self.v, value=2, command=self.manualname)
        self.R2.place(x=550, y=550)

        # progress bar
        self.bar = Progressbar(self.window, orient=HORIZONTAL, length=300)
        self.bar.place(x=200, y=500)
        l4 = Label(self.window, text="Progress:")
        l4.place(x=140, y=500)

        # run button
        self.b1 = Button(self.window, text="Run", width=25, state='normal', command=self.proteincompare)  # run button
        self.b1.place(x=700, y=525)

        if (self.v.get() == 1):
            self.e3.config(state='disabled')  # need it for initialise disable
            self.outText.set(self.name)
        elif (self.v.get() == 2):
            self.outText.set(self.e3.get())
        else:
            self.e3.insert(0, "Error")

    def defname(self):
        """
            set the output text to a default name and disable the entery for the name
        """
        self.e3.delete(0, END)
        self.e3.insert(0, self.name)
        self.e3.config(state='disabled')

    def manualname(self):
        """
            set the output text to the entery 3 and enable it
        """
        self.e3.config(state='normal')
        self.e3.delete(0, END)

    def mesagefinish(self):  # mesage box and call the function to clean the input info
        """
            Creater a message boc indicating the comparation has ended
        """
        messagebox.showinfo("Sucess", "Comparation complete")
        self.cleanscrean()

    def cleanscrean(self):  # clan all the inputs
        """
        clear all inputs and create a new default name
        """
        self.e1.delete(0, END)
        self.e2.delete(0, END)
        self.bar['value'] = 0
        self.v.set(1)
        self.name = datetime.now().strftime('%Y-%m-%d_%H-%M-%S') + "_output"
        self.outText.set(self.name)
        self.defname()

    def proteincompare(self):  # function to compare strings update bar and call msg finish
        """
        here we compare the proteins and update the task bar. Before however we will check the proteins are ready to
        be compared. We will perform some error checks
        """

        unaccepted_char = ['B', 'J', 'O', 'U', 'X', 'Z']        # Values not in Amino Acid Table

        for i in unaccepted_char:                               #Iterate through string to find if there is unaccepted char
            x = i in self.String1.get().upper()
            y = i in self.String2.get().upper()
            if (x == True or y == True):
                messagebox.showwarning('Entry Warning', 'Alphabetical value not part of Amino Acid dataset')
                break

        if (self.String1.get().isalpha() == False or self.String2.get().isalpha() == False):
            messagebox.showwarning("Entry Warning", "Warning special characters not accepted")


        if (self.e1.get() == "" or self.e2.get() == ""):
            messagebox.showwarning("Entry Warning", "Warning, empty selection. Please ensure BOTH forms are filled.")

        # IF no errors found we con proceed to compare the strings

        else:
            a = Compear(self.String1, self.String2, self.outText.get())
            self.bar['value'] += 50
            self.window.update_idletasks()
            a.extend_final_idx()
            x = 5
            while (x < 10):
                self.bar['value'] += 10
                time.sleep(1)
                x += 1
                self.window.update_idletasks()
            self.mesagefinish()



    # def chektext(self, event):
    #     """
    #         Check if we have text to enable or disable the button run
    #     """
    #     if (self.e1.get() != "" or self.e2.get() != ""):
    #         self.b1.config(state='normal')
    #     # elif (self.e1.get() == "" and self.e2.get() == ""):
    #     #     self.b1.config(state='disable')




    def startgui(self):
        """
            Needed to run the GUI
        """
        self.window.mainloop()

    def printinfo(self):
        """
        Print info use for testing
        """
        print("Output name : " + str(self.outText.get()) + "\n")
        print("Default Name: " + str(self.name) + "\n")
        print("String1     : " + str(self.String1.get()) + "\n")
        print("String2     : " + str(self.String2.get()) + "\n")


    def string_preper(self):
        """
        Prepare the strings to be compared.
        We can check for errors or characters that are not allowed
        Returns
        -------
        None.
        """
        # Give Error If....
        # entered strings have symbols
        # entered strings have non-protein valued letters
        # selection is empty


        # Except Letters.. these are NOT acronyms for proteins
        unaccepted = ['B', 'J', 'O', 'U', 'X', 'Z']

        self.String1.get().upper()
        self.String2.get().upper()

        # if (self.String1.get().isalpha() == False or self.String2.get().isalpha() == False):
        #     messagebox.showwarning("Entry Warning", "Warning Special Characters not Excepted")
        #     exit(1)


        # display warning that text not only contain letters

        for i in unaccepted:
            x = i in self.String1
            y = i in self.String2
            if (x == True or y == True):
                messagebox.showerror('Entry Error', 'Error: Symbols are not accepted as input!')
                exit(1)


a = GUI()
a.startgui()
