import sys

import docx
from docx.enum.text import WD_COLOR_INDEX

a = "THIS IS DUMMY DATA, TO TEST HIGHLIGHTING FOUND WORDS"
check_list = ["THIS", "DATA", "FOUND"]
# create an instance of a word doc
doc = docx.Document()
# add heading
doc.add_heading('Test Comparison Document', 0)
# add paragraphs
para = doc.add_paragraph(
    "Original Sequence: " + a
)
para0 = doc.add_paragraph(
    "Comparison Sequence: " + check_list[0] + " " + check_list[1] + " " + check_list[2]
)
para1 = doc.add_paragraph(
    "The following sequence will have the similarities highlighted in pink"
)
sim_para = doc.add_paragraph()
# what to look for

i = 0
result = dict()
for ele in check_list:
    if ele in a:
        start = a.index(ele)
        end = start + len(ele)
        sim_para.add_run(
            a[i:start]
        )
        sim_para.add_run(
            a[start:end]
        ).font.highlight_color = WD_COLOR_INDEX.PINK
        i = i + end
if end != sys.getsizeof(a):
    sim_para.add_run(
        a[end: (sys.getsizeof(a)-1)]
    )
# save document
doc.save('test2.docx')
